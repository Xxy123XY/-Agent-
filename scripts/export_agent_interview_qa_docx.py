from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "agent_interview_qa.md"
OUTPUT_DOCX = ROOT / "docs" / "agent_interview_qa_print.docx"


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    title = styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(18)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    heading1 = styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(13)
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(10)
    heading1.paragraph_format.space_after = Pt(6)
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    list_paragraph = styles["List Paragraph"]
    list_paragraph.font.name = "Times New Roman"
    list_paragraph.font.size = Pt(11)
    list_paragraph.paragraph_format.left_indent = Cm(0.74)
    list_paragraph.paragraph_format.first_line_indent = Cm(0)
    list_paragraph._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = footer.add_run()
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(10.5)

    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._element.append(fld_char1)
    run._element.append(instr_text)
    run._element.append(fld_char2)


def add_text_with_code(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(11)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")


def export() -> Path:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    title_text = "求职辅助多智能体 Agent 系统面试问答"
    for line in lines:
        if line.startswith("# "):
            title_text = line[2:].strip()
            break

    title = doc.add_paragraph(style="Title")
    title_run = title.add_run(title_text)
    set_east_asia_font(title_run, "黑体")
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("打印复习版")
    set_east_asia_font(subtitle_run, "宋体")
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.italic = True

    for raw_line in lines:
        line = raw_line.rstrip()

        if not line or line == "---" or line.startswith("# "):
            continue

        if line.startswith("## "):
            question = doc.add_paragraph(style="Heading 1")
            question.paragraph_format.keep_with_next = True
            run = question.add_run(line[3:].strip())
            set_east_asia_font(run, "黑体")
            run.font.size = Pt(13)
            run.font.bold = True
            continue

        if re.match(r"^\d+\.\s", line):
            paragraph = doc.add_paragraph(style="List Paragraph")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0.74)
            add_text_with_code(paragraph, line)
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Paragraph")
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_text_with_code(paragraph, f"• {line[2:].strip()}")
            continue

        paragraph = doc.add_paragraph(style="Normal")
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_text_with_code(paragraph, line.replace("  ", " ").strip())

    # Add a final section property touch so the page number is kept on all pages.
    doc.sections[0].start_type = WD_SECTION.NEW_PAGE
    add_page_number(doc.sections[0])

    target = _next_available_output_path(OUTPUT_DOCX)
    doc.save(target)
    return target


def _next_available_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    try:
        with open(path, "ab"):
            return path
    except PermissionError:
        pass

    stem = path.stem
    suffix = path.suffix
    parent = path.parent
    index = 2
    while True:
        candidate = parent / f"{stem}_v{index}{suffix}"
        if not candidate.exists():
            return candidate
        try:
            with open(candidate, "ab"):
                return candidate
        except PermissionError:
            index += 1


if __name__ == "__main__":
    path = export()
    print(path)
